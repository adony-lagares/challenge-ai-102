"""Azure Translator Service - Clean and functional implementation."""

import os
import uuid
import logging
from typing import Optional
from pathlib import Path

import requests
from docx import Document


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class AzureTranslator:
    """Azure Translator API client."""
    
    def __init__(
        self,
        subscription_key: Optional[str] = None,
        endpoint: str = "https://api.cognitive.microsofttranslator.com",
        location: str = "eastus2"
    ):
        self.subscription_key = subscription_key or os.getenv('AZURE_TRANSLATOR_KEY')
        self.endpoint = endpoint
        self.location = location
        
        if not self.subscription_key or self.subscription_key == 'YOUR_AZURE_SUBSCRIPTION_KEY':
            raise ValueError("Set AZURE_TRANSLATOR_KEY environment variable")
    
    def translate_text(
        self,
        text: str,
        target_language: str,
        source_language: str = 'en'
    ) -> str:
        """Translate text to target language."""
        if not text.strip():
            return text
        
        url = f"{self.endpoint}/translate"
        
        headers = {
            'Ocp-Apim-Subscription-Key': self.subscription_key,
            'Ocp-Apim-Subscription-Region': self.location,
            'Content-type': 'application/json',
            'X-ClientTraceId': str(uuid.uuid4())
        }
        
        params = {
            'api-version': '3.0',
            'from': source_language,
            'to': target_language
        }
        
        body = [{'text': text}]
        
        try:
            response = requests.post(url, params=params, headers=headers, json=body, timeout=30)
            response.raise_for_status()
            result = response.json()
            return result[0]['translations'][0]['text']
        except Exception as e:
            logger.error(f"Translation failed: {e}")
            raise
    
    def translate_document(
        self,
        document_path: str,
        target_language: str,
        source_language: str = 'en',
        output_path: Optional[str] = None
    ) -> str:
        """Translate Word document."""
        doc_path = Path(document_path)
        
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {document_path}")
        
        original_doc = Document(document_path)
        translated_doc = Document()
        
        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():
                translated_text = self.translate_text(
                    paragraph.text,
                    target_language,
                    source_language
                )
                translated_doc.add_paragraph(translated_text)
            else:
                translated_doc.add_paragraph()
        
        if not output_path:
            output_path = str(doc_path.parent / f"{doc_path.stem}_{target_language}.docx")
        
        translated_doc.save(output_path)
        logger.info(f"Document saved: {output_path}")
        return output_path


if __name__ == "__main__":
    translator = AzureTranslator(subscription_key="YOUR_AZURE_SUBSCRIPTION_KEY")
    
    # Example: Translate text
    result = translator.translate_text("Hello, world!", "pt-br")
    print(f"Translated: {result}")
    
    # Example: Translate document
    # translator.translate_document("example.docx", "pt-br")
